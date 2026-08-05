"""Create the first supervisor-facing research-notebook entry from the local template."""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(f"• {item}")


def add_section(document: Document, title: str, paragraphs: list[str] | None = None, bullets: list[str] | None = None) -> None:
    document.add_heading(title, level=1)
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)
    if bullets:
        add_bullets(document, bullets)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", default="docs/Research Notebook [template].docx")
    parser.add_argument("--output", default="docs/research_notebook/2026-08-05_entry_01.docx")
    args = parser.parse_args()

    document = Document(args.template)
    body = document._element.body
    for child in list(body):
        if child.tag.rsplit("}", 1)[-1] != "sectPr":
            body.remove(child)

    title = document.add_heading("Research Notebook", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Entry 1: Research foundation, English corpus, and publication")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=4, cols=2)
    metadata = [
        ("Date", "05 August 2026"),
        ("Researcher", "Ashioya Jotham"),
        ("Status", "In progress"),
        ("Project", "Runtime Safety Governors: Activation Steering as a Control API for Deployed Language Models"),
    ]
    for row, (key, value) in zip(table.rows, metadata):
        row.cells[0].text = key
        row.cells[1].text = value

    add_section(document, "Goal for today", [
        "Establish a reproducible first activation-steering experiment, prepare the approved English contrastive corpus, document the work clearly, and publish the research repository."
    ])
    add_section(document, "Context / Purpose", [
        "The project tests whether activation steering can provide a runtime safety intervention without retraining. Before measuring safety effects, the project needs a reproducible model path, provenance-linked contrastive data, and a clear record of decisions and limitations."
    ])
    add_section(document, "Method / How you did it", [
        "Reviewed the project proposal and built a configuration-based Python workflow for data validation, residual-activation capture, vector extraction, steering hooks, and Control Tax metrics. Used TransformerLens with GPT-2 Small for the first end-to-end smoke run. Used the local prior-work dataset, JailbreakBench, and harmless synthetic instruction pairs to prepare the English corpus."
    ])
    add_section(document, "What Happened (Step-by-step) / Experiment", bullets=[
        "Read the project proposal and aligned the repository to its research questions; removed the earlier product-gateway framing.",
        "Created the experiment package, configurations, validation scripts, curation workflow, tests, and research documentation.",
        "Installed TransformerLens and ran GPT-2 Small on CPU. Captured layer-0 residual activations from the fixture data, wrote a run manifest, and extracted a difference-in-means vector.",
        "Prepared 100 deceptive-reasoning pairs from the local CC-BY prior-work source, 100 harmless instruction-following/evasion pairs, and 100 restricted harmful-compliance pairs from JailbreakBench.",
        "Detected duplicate and degenerate source prompts during import. Preserved them as local audit material, excluded them from the public repository, and regenerated a structurally valid deceptive-reasoning set.",
        "Updated the README to describe the research questions, protocol, limitations, data controls, and phase sequence. Published the clean repository to GitHub."
    ])
    add_section(document, "Observations / Notes", bullets=[
        "The first GPT-2 vector is a plumbing check only: the fixture contains one safe/unsafe pair, so its bootstrap stability score is not scientific evidence.",
        "The 300 English pairs are structurally validated and marked approved following project-lead confirmation. Their provenance and reviewer-status fields are retained.",
        "Harmful-compliance source material is kept in the Git-ignored restricted data folder. It is not included in public commits or routine logs.",
        "This machine has no CUDA device. GPT-2 Small was suitable for the CPU smoke run; later Llama-3-8B sweeps require a GPU environment."
    ])
    add_section(document, "Results", bullets=[
        "Completed the first end-to-end capture -> manifest -> vector pipeline run on GPT-2 Small.",
        "Validated 300 approved English contrastive pairs (600 records): 100 pairs per target behavior.",
        "All local automated checks pass: 7 tests.",
        "Published commit 47be39b to the public GitHub repository."
    ])
    add_section(document, "Decisions & Why", bullets=[
        "Use GPT-2 Small first to validate the method quickly and cheaply before scaling to Llama-3-8B.",
        "Keep restricted harmful data and generated activation artifacts out of version control to reduce unnecessary exposure.",
        "Do not make a safety claim from the smoke artifact; proceed only after captures from the approved corpus and predeclared evaluation runs.",
        "Translate a frozen English evaluation subset next so that English-Swahili comparisons use matched meaning rather than unrelated prompts."
    ])
    add_section(document, "Next Actions", bullets=[
        "Researcher: select the frozen English evaluation subset and produce Swahili translations with bilingual review.",
        "Researcher: add translation provenance and quality notes, then validate that translated pairs remain in their original split.",
        "Researcher: run English/Swahili residual-stream captures and compute layer-wise representation similarity.",
        "Researcher: move the full extraction and Control Tax sweep to a GPU environment after the translation gate."
    ])
    add_section(document, "Time & Metadata", bullets=[
        "Time spent: not recorded for this retrospective entry; add actual hours before sending if required.",
        "Files updated: README.md; safety_governor/; scripts/; configs/; datasets/; docs/; tests/.",
        "Tags: #activation-steering #mechanistic-interpretability #dataset-curation #gpt2 #multilingual-safety",
        "Git commit: 47be39b — Initialize Runtime Safety Governor research pipeline."
    ])
    add_section(document, "Summary (short report)", [
        "Built and validated the first reproducible GPT-2 activation-capture and vector-extraction path, prepared a 300-pair approved English corpus, and published the research repository. The next step is a frozen English-to-Swahili evaluation subset with bilingual review; no safety conclusion is drawn from the one-pair smoke artifact."
    ])

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
